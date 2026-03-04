"""
Country PIM Transparency API.

Endpoints:
  POST /api/country-transparency        — generate a PIM transparency briefing (SSE)
  POST /api/country-transparency/export  — export generated content as DOCX or PDF
"""

from __future__ import annotations

import io
import json
import logging
import re

from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse, Response
from pydantic import BaseModel, Field

from app.config import settings
from app.generation.transparency_prompt import (
    format_transparency_records_context,
    get_country_transparency_prompt,
)
from app.generation.prompts import format_documents
from app.ingestion.repo_source import fetch_records_with_docs, resolve_country_name
from app.retrieval.retriever import get_retriever

logger = logging.getLogger(__name__)
router = APIRouter(tags=["country-transparency"])


# ── Models ────────────────────────────────────────────────────

class TransparencyRequest(BaseModel):
    country_iso3: str = Field(description="ISO3 country code")


class TransparencyExportRequest(BaseModel):
    content: str = Field(description="The generated markdown content")
    country_name: str = Field(description="Country name for the document title")
    format: str = Field(description="'docx' or 'pdf'")


# ── Helpers ───────────────────────────────────────────────────

def _sse(data: dict) -> str:
    return f"data: {json.dumps(data)}\n\n"


def _get_transparency_llm():
    """Return a ChatAnthropic instance with higher max_tokens for briefing generation."""
    from langchain_anthropic import ChatAnthropic

    return ChatAnthropic(
        model=settings.llm_model,
        temperature=settings.llm_temperature,
        max_tokens=8192,
        anthropic_api_key=settings.anthropic_api_key,
    )


# ── POST /api/country-transparency ───────────────────────────

@router.post("/country-transparency")
async def generate_country_transparency(request: TransparencyRequest):
    """Generate a PIM transparency briefing and stream via SSE."""
    return StreamingResponse(
        _transparency_stream(request.country_iso3),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


async def _transparency_stream(country_iso3: str):
    """Generator that builds context, calls the LLM, and yields SSE events."""
    try:
        # 1. Resolve country name
        country_name = resolve_country_name(country_iso3) or country_iso3
        yield _sse({"type": "status", "data": f"Preparing transparency briefing for {country_name}..."})

        # 2. Fetch structured policy records
        records = fetch_records_with_docs(country=country_iso3)
        policy_context = format_transparency_records_context(records)

        yield _sse({
            "type": "status",
            "data": f"Found {len(records)} policy records. Retrieving document context...",
        })

        # 3. Retrieve RAG document chunks for this country
        try:
            retriever = get_retriever(filters={"country": country_iso3}, k=20)
            docs = retriever.invoke(
                f"Public investment management institutional framework policy transparency for {country_name}"
            )
            rag_context = format_documents(docs) if docs else "No indexed documents found for this country."
            yield _sse({
                "type": "status",
                "data": f"Retrieved {len(docs)} document excerpts. Generating transparency briefing...",
            })
        except Exception as e:
            logger.warning("RAG retrieval failed for %s: %s", country_iso3, e)
            rag_context = "No indexed documents available."
            yield _sse({
                "type": "status",
                "data": "No indexed documents available. Generating from general knowledge...",
            })

        # 4. Build the prompt and stream the LLM response
        prompt = get_country_transparency_prompt()
        llm = _get_transparency_llm()
        chain = prompt | llm

        async for chunk in chain.astream({
            "country_name": country_name,
            "policy_records_context": policy_context,
            "rag_context": rag_context,
        }):
            if chunk.content:
                yield _sse({"type": "token", "data": chunk.content})

        yield _sse({"type": "done"})

    except Exception as e:
        logger.exception("Transparency briefing generation failed for %s", country_iso3)
        yield _sse({"type": "error", "data": str(e)})


# ── POST /api/country-transparency/export ─────────────────────

@router.post("/country-transparency/export")
async def export_country_transparency(request: TransparencyExportRequest):
    """Export a generated transparency briefing as DOCX or PDF."""
    if request.format == "docx":
        return _export_docx(request.content, request.country_name)
    elif request.format == "pdf":
        return _export_pdf(request.content, request.country_name)
    else:
        raise HTTPException(status_code=400, detail="Format must be 'docx' or 'pdf'")


def _export_docx(content: str, country_name: str) -> Response:
    """Generate a Word document from markdown content."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()

    # Page setup (A4)
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Title
    title = doc.add_heading(f"PIM Transparency Briefing: {country_name}", level=1)
    title.runs[0].font.size = Pt(16)
    title.runs[0].font.color.rgb = RGBColor(0x0e, 0x7c, 0x47)  # Green accent

    # Parse markdown content into sections
    lines = content.split("\n")
    i = 0
    in_table = False
    table_rows: list[list[str]] = []

    while i < len(lines):
        line = lines[i]

        # Headings
        if line.startswith("## "):
            if in_table and table_rows:
                _add_table_to_doc(doc, table_rows)
                table_rows = []
                in_table = False
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            if in_table and table_rows:
                _add_table_to_doc(doc, table_rows)
                table_rows = []
                in_table = False
            doc.add_heading(line[4:].strip(), level=3)

        # Table rows
        elif line.strip().startswith("|") and line.strip().endswith("|"):
            # Skip separator rows (|---|---|)
            if re.match(r"^\|[\s\-:|]+\|$", line.strip()):
                i += 1
                continue
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            table_rows.append(cells)
            in_table = True

        # Bullet points
        elif line.strip().startswith("- ") or line.strip().startswith("* "):
            if in_table and table_rows:
                _add_table_to_doc(doc, table_rows)
                table_rows = []
                in_table = False
            text = line.strip().lstrip("-* ").strip()
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_run(p, text)

        # Bold-only lines (like **text**)
        elif line.strip().startswith("**") and line.strip().endswith("**"):
            if in_table and table_rows:
                _add_table_to_doc(doc, table_rows)
                table_rows = []
                in_table = False
            p = doc.add_paragraph()
            run = p.add_run(line.strip().strip("*"))
            run.bold = True

        # Regular paragraphs
        elif line.strip():
            if in_table and table_rows:
                _add_table_to_doc(doc, table_rows)
                table_rows = []
                in_table = False
            p = doc.add_paragraph()
            _add_formatted_run(p, line.strip())

        else:
            # Empty line — flush any pending table
            if in_table and table_rows:
                _add_table_to_doc(doc, table_rows)
                table_rows = []
                in_table = False

        i += 1

    # Flush final table
    if table_rows:
        _add_table_to_doc(doc, table_rows)

    # Save to buffer
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"PIM_Transparency_{country_name.replace(' ', '_')}.docx"
    return Response(
        content=buf.read(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _add_table_to_doc(doc, rows: list[list[str]]):
    """Add a markdown-parsed table to the Word document."""
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn

    if not rows:
        return

    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = "Table Grid"

    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < num_cols:
                cell = table.rows[row_idx].cells[col_idx]
                cell.text = ""
                p = cell.paragraphs[0]
                # Strip markdown bold markers
                clean_text = re.sub(r"\*\*(.*?)\*\*", r"\1", cell_text)
                run = p.add_run(clean_text)
                run.font.size = Pt(9)

                # Style header row — green for transparency theme
                if row_idx == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    shading = cell._element.get_or_add_tcPr()
                    bg = shading.makeelement(
                        qn("w:shd"),
                        {
                            qn("w:val"): "clear",
                            qn("w:color"): "auto",
                            qn("w:fill"): "0e7c47",
                        },
                    )
                    shading.append(bg)

    doc.add_paragraph()  # Spacing after table


def _add_formatted_run(paragraph, text: str):
    """Add text to a paragraph, handling **bold** markers."""
    from docx.shared import Pt

    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(10)
        else:
            run = paragraph.add_run(part)
            run.font.size = Pt(10)


def _export_pdf(content: str, country_name: str) -> Response:
    """Generate a PDF document from markdown content."""
    from fpdf import FPDF

    pdf = FPDF(orientation="P", unit="mm", format="A4")
    pdf.set_auto_page_break(auto=True, margin=25)
    pdf.add_page()

    # Header
    pdf.set_font("Helvetica", "I", 8)
    pdf.set_text_color(128, 128, 128)
    pdf.cell(0, 5, f"PIM Transparency Briefing — {country_name}", align="R", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(5)

    # Title — green accent
    pdf.set_font("Helvetica", "B", 16)
    pdf.set_text_color(14, 124, 71)
    pdf.cell(0, 10, f"PIM Transparency Briefing: {country_name}", new_x="LMARGIN", new_y="NEXT")
    pdf.set_draw_color(14, 124, 71)
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(8)
    pdf.set_text_color(0, 0, 0)

    lines = content.split("\n")
    i = 0
    table_rows: list[list[str]] = []
    in_table = False

    while i < len(lines):
        line = lines[i]

        # Headings
        if line.startswith("## "):
            if in_table and table_rows:
                _add_table_to_pdf(pdf, table_rows)
                table_rows = []
                in_table = False
            pdf.ln(4)
            pdf.set_font("Helvetica", "B", 12)
            pdf.cell(0, 7, line[3:].strip(), new_x="LMARGIN", new_y="NEXT")
            pdf.ln(2)
        elif line.startswith("### "):
            if in_table and table_rows:
                _add_table_to_pdf(pdf, table_rows)
                table_rows = []
                in_table = False
            pdf.ln(2)
            pdf.set_font("Helvetica", "B", 10)
            pdf.cell(0, 6, line[4:].strip(), new_x="LMARGIN", new_y="NEXT")
            pdf.ln(1)

        # Table rows
        elif line.strip().startswith("|") and line.strip().endswith("|"):
            if re.match(r"^\|[\s\-:|]+\|$", line.strip()):
                i += 1
                continue
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            table_rows.append(cells)
            in_table = True

        # Bullet points
        elif line.strip().startswith("- ") or line.strip().startswith("* "):
            if in_table and table_rows:
                _add_table_to_pdf(pdf, table_rows)
                table_rows = []
                in_table = False
            text = re.sub(r"\*\*(.*?)\*\*", r"\1", line.strip().lstrip("-* ").strip())
            pdf.set_font("Helvetica", "", 9)
            pdf.cell(5, 5, chr(8226))
            pdf.multi_cell(0, 5, f" {text}", new_x="LMARGIN", new_y="NEXT")

        # Regular text
        elif line.strip():
            if in_table and table_rows:
                _add_table_to_pdf(pdf, table_rows)
                table_rows = []
                in_table = False
            text = re.sub(r"\*\*(.*?)\*\*", r"\1", line.strip())
            pdf.set_font("Helvetica", "", 9)
            pdf.multi_cell(0, 5, text, new_x="LMARGIN", new_y="NEXT")

        else:
            if in_table and table_rows:
                _add_table_to_pdf(pdf, table_rows)
                table_rows = []
                in_table = False
            pdf.ln(2)

        i += 1

    if table_rows:
        _add_table_to_pdf(pdf, table_rows)

    # Footer — page numbers
    pdf.set_y(-15)
    pdf.set_font("Helvetica", "", 8)
    pdf.cell(0, 10, f"Page {pdf.page_no()}", align="C")

    buf = io.BytesIO()
    pdf.output(buf)
    buf.seek(0)

    filename = f"PIM_Transparency_{country_name.replace(' ', '_')}.pdf"
    return Response(
        content=buf.read(),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _add_table_to_pdf(pdf, rows: list[list[str]]):
    """Add a table to the PDF."""
    if not rows:
        return

    num_cols = max(len(r) for r in rows)
    page_width = pdf.w - pdf.l_margin - pdf.r_margin
    col_width = page_width / num_cols

    for row_idx, row_data in enumerate(rows):
        # Header row styling — green for transparency
        if row_idx == 0:
            pdf.set_fill_color(14, 124, 71)
            pdf.set_text_color(255, 255, 255)
            pdf.set_font("Helvetica", "B", 8)
        else:
            pdf.set_fill_color(255, 255, 255)
            pdf.set_text_color(0, 0, 0)
            pdf.set_font("Helvetica", "", 8)

        row_height = 6
        # Calculate max height needed for this row
        for col_idx in range(num_cols):
            cell_text = row_data[col_idx] if col_idx < len(row_data) else ""
            clean = re.sub(r"\*\*(.*?)\*\*", r"\1", cell_text)
            # Approximate needed lines
            lines_needed = max(1, len(clean) // int(col_width / 2) + 1)
            row_height = max(row_height, lines_needed * 5)

        x_start = pdf.get_x()
        y_start = pdf.get_y()

        # Check if we need a new page
        if y_start + row_height > pdf.h - pdf.b_margin:
            pdf.add_page()
            y_start = pdf.get_y()

        for col_idx in range(num_cols):
            cell_text = row_data[col_idx] if col_idx < len(row_data) else ""
            clean = re.sub(r"\*\*(.*?)\*\*", r"\1", cell_text)
            pdf.set_xy(x_start + col_idx * col_width, y_start)
            pdf.multi_cell(
                col_width, 5, clean,
                border=1,
                fill=(row_idx == 0),
                new_x="RIGHT", new_y="TOP",
            )

        pdf.set_xy(x_start, y_start + row_height)

    pdf.ln(4)
