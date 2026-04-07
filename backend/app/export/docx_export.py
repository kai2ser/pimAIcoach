"""
Shared DOCX export — generate a Word document from markdown content.

Parameterized by a `theme` dict so both the Profile (blue) and Transparency
(green) endpoints can reuse the same logic with different colours and labels.
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass, field

from fastapi.responses import Response


# ── Theme configuration ──────────────────────────────────────────────────────

@dataclass(frozen=True)
class DocxTheme:
    """Colour / label overrides for a DOCX export."""

    title_prefix: str = "Public Investment Management Context"
    filename_prefix: str = "PIM_Profile"
    # RGB tuples for the accent colour used in title & table headers
    accent_rgb: tuple[int, int, int] = (0x1D, 0x4E, 0xD8)   # blue
    # Hex string (without #) for table-header shading
    header_fill_hex: str = "1d4ed8"


PROFILE_THEME = DocxTheme()

TRANSPARENCY_THEME = DocxTheme(
    title_prefix="PIM Transparency Briefing",
    filename_prefix="PIM_Transparency",
    accent_rgb=(0x0E, 0x7C, 0x47),
    header_fill_hex="0e7c47",
)


# ── Public API ───────────────────────────────────────────────────────────────

def export_docx(content: str, country_name: str, theme: DocxTheme = PROFILE_THEME) -> Response:
    """Generate a Word document from markdown *content* and return a FastAPI Response."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor

    doc = Document()

    # Page setup (A4)
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Title
    title = doc.add_heading(f"{theme.title_prefix}: {country_name}", level=1)
    title.runs[0].font.size = Pt(16)
    title.runs[0].font.color.rgb = RGBColor(*theme.accent_rgb)

    # Parse markdown content
    lines = content.split("\n")
    i = 0
    in_table = False
    table_rows: list[list[str]] = []

    while i < len(lines):
        line = lines[i]

        if line.startswith("## "):
            if in_table and table_rows:
                _add_table_to_doc(doc, table_rows, theme)
                table_rows, in_table = [], False
            doc.add_heading(line[3:].strip(), level=2)

        elif line.startswith("### "):
            if in_table and table_rows:
                _add_table_to_doc(doc, table_rows, theme)
                table_rows, in_table = [], False
            doc.add_heading(line[4:].strip(), level=3)

        elif line.strip().startswith("|") and line.strip().endswith("|"):
            if re.match(r"^\|[\s\-:|]+\|$", line.strip()):
                i += 1
                continue
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            table_rows.append(cells)
            in_table = True

        elif line.strip().startswith("- ") or line.strip().startswith("* "):
            if in_table and table_rows:
                _add_table_to_doc(doc, table_rows, theme)
                table_rows, in_table = [], False
            text = line.strip().lstrip("-* ").strip()
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_run(p, text)

        elif line.strip().startswith("**") and line.strip().endswith("**"):
            if in_table and table_rows:
                _add_table_to_doc(doc, table_rows, theme)
                table_rows, in_table = [], False
            p = doc.add_paragraph()
            run = p.add_run(line.strip().strip("*"))
            run.bold = True

        elif line.strip():
            if in_table and table_rows:
                _add_table_to_doc(doc, table_rows, theme)
                table_rows, in_table = [], False
            p = doc.add_paragraph()
            _add_formatted_run(p, line.strip())

        else:
            if in_table and table_rows:
                _add_table_to_doc(doc, table_rows, theme)
                table_rows, in_table = [], False

        i += 1

    if table_rows:
        _add_table_to_doc(doc, table_rows, theme)

    # Save to buffer
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"{theme.filename_prefix}_{country_name.replace(' ', '_')}.docx"
    return Response(
        content=buf.read(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ── Helpers ──────────────────────────────────────────────────────────────────

def _add_table_to_doc(doc, rows: list[list[str]], theme: DocxTheme):
    """Add a markdown-parsed table to the Word document."""
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn

    if not rows:
        return

    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = "Table Grid"

    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < num_cols:
                cell = table.rows[row_idx].cells[col_idx]
                cell.text = ""
                p = cell.paragraphs[0]
                clean_text = re.sub(r"\*\*(.*?)\*\*", r"\1", cell_text)
                run = p.add_run(clean_text)
                run.font.size = Pt(9)

                if row_idx == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    shading = cell._element.get_or_add_tcPr()
                    bg = shading.makeelement(
                        qn("w:shd"),
                        {
                            qn("w:val"): "clear",
                            qn("w:color"): "auto",
                            qn("w:fill"): theme.header_fill_hex,
                        },
                    )
                    shading.append(bg)

    doc.add_paragraph()


def _add_formatted_run(paragraph, text: str):
    """Add text to a paragraph, handling **bold** markers."""
    from docx.shared import Pt

    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(10)
        else:
            run = paragraph.add_run(part)
            run.font.size = Pt(10)
